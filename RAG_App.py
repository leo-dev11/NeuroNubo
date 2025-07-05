import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from typing import List
import shutil
from pydantic import BaseModel
from langchain.chat_models import ChatOpenAI
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document
from langchain.chains import RetrievalQA
from fastapi.middleware.cors import CORSMiddleware
from langchain.prompts import PromptTemplate
from langgraph.graph import StateGraph
from langchain.tools import Tool
from langchain.agents import initialize_agent, AgentType
from langchain.prompts.chat import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
from dotenv import load_dotenv

# ---- FastAPI setup ----
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "./documents"
#Rutas del vectorstore
INDEX_DIR = "faiss_index"
FAISS_INDEX_FILE = os.path.join(INDEX_DIR, "index.faiss")

os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}

class QueryRequest(BaseModel):
    question: str

load_dotenv()# Carga las variables desde .env

api_key = os.getenv("OPENROUTER_API_KEY")
api_base = os.getenv("OPENROUTER_BASE_URL")

# ---- Leer documentos desde carpeta ----
def read_txt(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()

def read_pdf(filepath):
    text = ""
    with fitz.open(filepath) as doc:
        for page in doc:
            text += page.get_text()
    return text

def read_docx(filepath):
    doc = DocxDocument(filepath)
    return "\n".join([para.text for para in doc.paragraphs])

def load_documents_from_folder(folder_path="documents"):
    supported = {".pdf": read_pdf, ".docx": read_docx, ".txt": read_txt}
    docs = []
    for filename in os.listdir(folder_path):
        ext = os.path.splitext(filename)[1].lower()
        if ext in supported:
            path = os.path.join(folder_path, filename)
            try:
                content = supported[ext](path)
                if content.strip():
                    docs.append(Document(page_content=content, metadata={"source": filename}))
            except Exception as e:
                print(f"Error leyendo {filename}: {e}")
    return docs

def extract_text_from_file(filepath, ext):
    if ext == ".pdf":
        return read_pdf(filepath)
    elif ext == ".docx":
        return read_docx(filepath)
    elif ext == ".txt":
        return read_txt(filepath)
    else:
        raise ValueError("Extensión no soportada.")
# ---- Crear documentos y embeddings ----
user_docs = load_documents_from_folder("documents")

splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
split_docs = splitter.split_documents(user_docs)

embeddings = HuggingFaceEmbeddings(model_name="Qwen/Qwen3-Embedding-0.6B")

if os.path.exists(FAISS_INDEX_FILE):
    print("✅ Cargando vectorstore existente...")
    vectorstore = FAISS.load_local(
        INDEX_DIR,
        embeddings,
        allow_dangerous_deserialization=True
    )
else:
    print("📦 No se encontró el índice, creando uno nuevo desde documentos...")
    vectorstore = FAISS.from_documents(split_docs, embeddings)
    vectorstore.save_local(INDEX_DIR)

llm = ChatOpenAI(
    model="google/gemini-2.5-flash-lite-preview-06-17",
    temperature=0.0,
    openai_api_key=api_key,
    openai_api_base=api_base
)

chat_prompt = ChatPromptTemplate.from_messages([
    SystemMessagePromptTemplate.from_template(
        "Eres un asistente experto y amigable. Usa el contexto y tu conocimiento general para responder. "
        "Si la respuesta no está en el contexto, responde igual usando tu conocimiento, sin decir que no estaba en el contexto."
    ),
    HumanMessagePromptTemplate.from_template("Pregunta:\n{question}\n\nContexto:\n{context}")
])

qa_chain = RetrievalQA.from_chain_type(
    llm=llm,
    retriever=vectorstore.as_retriever(search_kwargs={"k": 5}),
    chain_type="stuff",
    chain_type_kwargs={"prompt": chat_prompt}
)
# ---- RAG Tool ----
rag_tool = Tool(
    name="document_search",
    func=lambda q: qa_chain.run(q),
    description="Usa esto para responder preguntas sobre documentos cargados por el usuario."
)
agent_executor = initialize_agent(
    tools=[rag_tool],
    llm=llm,
    agent=AgentType.OPENAI_FUNCTIONS,  # Usa Function Calling
    verbose=True
)

# ---- LangGraph setup ----
class RAGState(dict):
    pass

def preprocess_question(state: dict) -> dict:
    question = state["question"].strip().lower()
    return RAGState({"question": question})

def generate_answer_from_qa_chain(state: dict) -> dict:
    question = state.get("question", "")
    try:
        result = qa_chain.invoke({"query": question})
        return RAGState({
            "question": question,
            "answer": result.get("result", ""),
            "context": result.get("source_documents", [])
        })
    except Exception as e:
        return RAGState({
            "question": question,
            "answer": f"Error: {str(e)}",
            "context": []
        })
def generate_answer_with_agent(state: dict) -> dict:
    question = state.get("question", "")
    try:
        answer = agent_executor.run(question)
        return RAGState({
            "question": question,
            "answer": answer,
        })
    except Exception as e:
        return RAGState({
            "question": question,
            "answer": f"Error: {str(e)}",
        })

builder = StateGraph(dict)

builder.add_node("preprocess_question", preprocess_question)
builder.add_node("generate_answer", generate_answer_from_qa_chain)

builder.set_entry_point("preprocess_question")
builder.add_edge("preprocess_question", "generate_answer")
builder.set_finish_point("generate_answer")

graph = builder.compile()

# ---- API endpoint ----
class QueryRequest(BaseModel):
    question: str

@app.post("/query")
async def query_agent(req: QueryRequest):
    resultado = graph.invoke({"question": req.question})
    return {"answer": resultado["answer"]}

# Correr en la terminal con "uvicorn RAG_App:app --reload

# POST para subir documentos
@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    file_ext = os.path.splitext(file.filename)[1].lower()

    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Tipo de archivo no permitido.")

    file_path = os.path.join(UPLOAD_DIR, file.filename)

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        text = extract_text_from_file(file_path, file_ext)
        if not text.strip():
            raise ValueError("El archivo está vacío o no contiene texto útil.")

        doc = Document(page_content=text, metadata={"source": file.filename})
        new_chunks = splitter.split_documents([doc])

        vectorstore.add_documents(new_chunks)

        vectorstore.save_local("faiss_index")

        return  {
            "filename": file.filename,
            "status": "Documento cargado e indexado correctamente.",
            "chunks": len(new_chunks)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al procesar e indexar el documento: {e}")

# GET para listar documentos cargados
@app.get("/documents")
async def list_documents():
    files = os.listdir(UPLOAD_DIR)
    allowed_files = [f for f in files if os.path.splitext(f)[1].lower() in ALLOWED_EXTENSIONS]
    return {"uploaded_documents": allowed_files}